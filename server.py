#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
MCP Docx Processing Service
Provides various operations for docx documents, including querying, adding, modifying, deleting, and font style settings
Implemented using the official MCP library
"""

import os
import sys
import argparse
import tempfile
import logging
import traceback
from contextlib import asynccontextmanager
from typing import AsyncIterator, Dict, Any, Optional

from mcp.server.fastmcp import FastMCP, Context
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Parse server mode flags early (before tool registration) so they're available at runtime.
# parse_known_args lets the MCP SDK handle its own argv untouched.
_arg_parser = argparse.ArgumentParser(add_help=False)
_arg_parser.add_argument(
    "--readonly",
    action="store_true",
    help="Disable all write operations (open and read-only tools still work)"
)
_arg_parser.add_argument(
    "--cwd-only",
    action="store_true",
    help="Restrict file access to the current working directory"
)
_server_args, _remaining_argv = _arg_parser.parse_known_args()
sys.argv = [sys.argv[0]] + _remaining_argv

READONLY_MODE: bool = _server_args.readonly
CWD_ONLY_MODE: bool = _server_args.cwd_only

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler(os.path.join(tempfile.gettempdir(), "docx_mcp_server.log")),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("DocxMCPServer")

# Create a state file for restoring state when MCP service restarts
CURRENT_DOC_FILE = os.path.join(tempfile.gettempdir(), "docx_mcp_current_doc.txt")

# ---------------------------------------------------------------------------
# Mode guards
# ---------------------------------------------------------------------------

def _check_readonly() -> Optional[str]:
    """Return an error string if the server is in read-only mode, else None."""
    if READONLY_MODE:
        return "Server is running in read-only mode: write operations are disabled"
    return None


def _check_path_allowed(path: str) -> Optional[str]:
    """Return an error string if *path* is outside the CWD when --cwd-only is active, else None."""
    if not CWD_ONLY_MODE:
        return None
    cwd_real = os.path.realpath(os.getcwd())
    path_real = os.path.realpath(os.path.abspath(path))
    # Accept the cwd itself or anything strictly inside it
    if path_real != cwd_real and not path_real.startswith(cwd_real + os.sep):
        return (
            f"Access denied: '{path}' is outside the allowed directory ({os.getcwd()}). "
            "Server was started with --cwd-only."
        )
    return None

class DocxProcessor:
    """Class for processing Docx documents, implementing various document operations"""
    
    def __init__(self):
        self.documents = {}  # Store opened documents
        self.current_document = None
        self.current_file_path = None
        
        # Try to load current document from state file
        self._load_current_document()
    
    def _load_current_document(self):
        """Load current document from state file"""
        if not os.path.exists(CURRENT_DOC_FILE):
            return False
        
        try:
            with open(CURRENT_DOC_FILE, 'r', encoding='utf-8') as f:
                file_path = f.read().strip()
            
            if file_path and os.path.exists(file_path):
                try:
                    self.current_file_path = file_path
                    self.current_document = Document(file_path)
                    self.documents[file_path] = self.current_document
                    return True
                except Exception as e:
                    logger.error(f"Failed to load document at {file_path}: {e}")
                    # Delete invalid state file to prevent future loading attempts
                    try:
                        os.remove(CURRENT_DOC_FILE)
                        logger.info(f"Removed invalid state file pointing to {file_path}")
                    except Exception as e_remove:
                        logger.error(f"Failed to remove state file: {e_remove}")
            else:
                # Delete invalid state file if path is empty or file doesn't exist
                try:
                    os.remove(CURRENT_DOC_FILE)
                    logger.info("Removed invalid state file with non-existent document path")
                except Exception as e_remove:
                    logger.error(f"Failed to remove state file: {e_remove}")
        except Exception as e:
            logger.error(f"Failed to load current document: {e}")
            # Delete corrupted state file
            try:
                os.remove(CURRENT_DOC_FILE)
                logger.info("Removed corrupted state file")
            except Exception as e_remove:
                logger.error(f"Failed to remove state file: {e_remove}")
        
        return False
    
    def _save_current_document(self):
        """Save current document path to state file"""
        if not self.current_file_path:
            return False
        
        try:
            with open(CURRENT_DOC_FILE, 'w', encoding='utf-8') as f:
                f.write(self.current_file_path)
            return True
        except Exception as e:
            logger.error(f"Failed to save current document path: {e}")
        
        return False
    
    def save_state(self):
        """Save processor state"""
        # Save current document
        if self.current_document and self.current_file_path:
            try:
                self.current_document.save(self.current_file_path)
                self._save_current_document()
            except Exception as e:
                logger.error(f"Failed to save current document: {e}")
    
    def load_state(self):
        """Load processor state"""
        self._load_current_document()

    # ... Keep all original document processing methods ...

# Create global processor instance
processor = DocxProcessor()

@asynccontextmanager
async def server_lifespan(server: FastMCP) -> AsyncIterator[Dict[str, Any]]:
    """Manage server lifecycle"""
    try:
        # Start server with clean state
        logger.info("DocxProcessor MCP server starting with clean state...")
        # Do not attempt to load any previous state
        yield {"processor": processor}
    finally:
        # Save state when server shuts down
        logger.info("DocxProcessor MCP server shutting down...")
        if processor.current_document and processor.current_file_path:
            processor.save_state()
        else:
            logger.info("No document open, not saving state")

# Create MCP server
mcp = FastMCP(
    name="DocxProcessor",
    instructions="Word document processing service, providing functions to create, edit, and query documents",
    lifespan=server_lifespan
)

@mcp.tool()
def create_document(ctx: Context, file_path: str) -> str:
    """
    Create a new Word document

    Parameters:
    - file_path: Document save path
    """
    try:
        if err := _check_readonly():
            return err
        if err := _check_path_allowed(file_path):
            return err
        processor.current_document = Document()
        processor.current_file_path = file_path
        processor.documents[file_path] = processor.current_document
        
        # Save document
        processor.current_document.save(file_path)
        
        return f"Document created successfully: {file_path}"
    except Exception as e:
        error_msg = f"Failed to create document: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def open_document(ctx: Context, file_path: str) -> str:
    """
    Open an existing Word document

    Parameters:
    - file_path: Path to the document to open
    """
    try:
        if err := _check_path_allowed(file_path):
            return err
        if not os.path.exists(file_path):
            return f"File does not exist: {file_path}"
        
        processor.current_document = Document(file_path)
        processor.current_file_path = file_path
        processor.documents[file_path] = processor.current_document
        
        return f"Document opened successfully: {file_path}"
    except Exception as e:
        error_msg = f"Failed to open document: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def save_document(ctx: Context) -> str:
    """
    Save the currently open Word document to the original file (update the original file)
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        if not processor.current_file_path:
            return "Current document has not been saved before, please use save_as_document to specify a save path"
            
        # Save to original file path
        processor.current_document.save(processor.current_file_path)
        
        return f"Document saved successfully to original file: {processor.current_file_path}"
    except Exception as e:
        error_msg = f"Failed to save document: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def add_paragraph(
    ctx: Context,
    text: str,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    font_size: Optional[int] = None,
    font_name: Optional[str] = None,
    color: Optional[str] = None,
    alignment: Optional[str] = None
) -> str:
    """
    Add paragraph text to document

    Parameters:
    - text: Paragraph text content
    - bold: Whether to bold
    - italic: Whether to italicize
    - underline: Whether to underline
    - font_size: Font size (points)
    - font_name: Font name
    - color: Text color (format: #FF0000)
    - alignment: Alignment (left, center, right, justify)
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        # Add paragraph
        paragraph = processor.current_document.add_paragraph(text)
        
        # Apply additional formatting
        if paragraph.runs:
            run = paragraph.runs[0]
            run.bold = bold
            run.italic = italic
            run.underline = underline
            
            # Set font size
            if font_size:
                run.font.size = Pt(font_size)
            
            # Set font name
            if font_name:
                run.font.name = font_name
                # Set East Asian font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            
            # Set font color
            if color and color.startswith('#') and len(color) == 7:
                r = int(color[1:3], 16)
                g = int(color[3:5], 16)
                b = int(color[5:7], 16)
                run.font.color.rgb = RGBColor(r, g, b)
        
        # Set alignment
        if alignment:
            if alignment == "left":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif alignment == "center":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif alignment == "right":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif alignment == "justify":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        return "Paragraph added"
    except Exception as e:
        error_msg = f"Failed to add paragraph: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def add_heading(ctx: Context, text: str, level: int) -> str:
    """
    Add heading to document

    Parameters:
    - text: Heading text
    - level: Heading level (1-9)
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        processor.current_document.add_heading(text, level=level)
        
        return f"Added level {level} heading"
    except Exception as e:
        error_msg = f"Failed to add heading: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def add_table(ctx: Context, rows: int, cols: int, data: Optional[list] = None) -> str:
    """
    Add table to document

    Parameters:
    - rows: Number of rows
    - cols: Number of columns
    - data: Table data, two-dimensional array
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        table = processor.current_document.add_table(rows=rows, cols=cols, style="Table Grid")
        
        # Fill table data
        if data:
            for i, row_data in enumerate(data):
                if i < rows:
                    row = table.rows[i]
                    for j, cell_text in enumerate(row_data):
                        if j < cols:
                            row.cells[j].text = str(cell_text)
        
        return f"Added {rows}x{cols} table"
    except Exception as e:
        error_msg = f"Failed to add table: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def get_document_info(ctx: Context) -> str:
    """
    Get document information, including paragraph count, table count, styles, etc.
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        # Get basic document information
        sections_count = len(doc.sections)
        paragraphs_count = len(doc.paragraphs)
        tables_count = len(doc.tables)
        
        # Get style list
        paragraph_styles = []
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                paragraph_styles.append(style.name)
        
        # Build information string
        info = f"Document path: {processor.current_file_path}\n"
        info += f"Section count: {sections_count}\n"
        info += f"Paragraph count: {paragraphs_count}\n"
        info += f"Table count: {tables_count}\n"
        info += f"Available paragraph styles: {', '.join(paragraph_styles[:10])}..."
        
        return info
    except Exception as e:
        error_msg = f"Failed to get document information: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def search_text(ctx: Context, keyword: str) -> str:
    """
    Search for text in the document
    
    Parameters:
    - keyword: Keyword to search for
    """
    try:
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        results = []
        
        # Search in paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            if keyword in paragraph.text:
                results.append({
                    "type": "paragraph",
                    "index": i,
                    "text": paragraph.text
                })
        
        # Search in tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if keyword in cell.text:
                        results.append({
                            "type": "table cell",
                            "table_index": t_idx,
                            "row": r_idx,
                            "column": c_idx,
                            "text": cell.text
                        })
        
        if not results:
            return f"Keyword '{keyword}' not found"
        
        # Build response
        response = f"Found {len(results)} occurrences of '{keyword}':\n\n"
        for idx, result in enumerate(results):
            response += f"{idx+1}. {result['type']} "
            if result['type'] == "paragraph":
                response += f"index {result['index']}: {result['text'][:100]}"
                if len(result['text']) > 100:
                    response += "..."
                response += "\n"
            else:
                response += f"in table {result['table_index']} at cell ({result['row']},{result['column']}): {result['text'][:100]}"
                if len(result['text']) > 100:
                    response += "..."
                response += "\n"
        
        return response
    except Exception as e:
        error_msg = f"Failed to search text: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def search_and_replace(ctx: Context, keyword: str, replace_with: str, preview_only: bool = False) -> str:
    """
    Search and replace text in the document, providing detailed replacement information and preview options

    Parameters:
    - keyword: Keyword to search for
    - replace_with: Text to replace with
    - preview_only: Whether to only preview without actually replacing, default is False
    """
    try:
        if not preview_only:
            if err := _check_readonly():
                return err
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        results = []
        
        # Search in paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            if keyword in paragraph.text:
                # Save original text and replaced text
                original_text = paragraph.text
                replaced_text = original_text.replace(keyword, replace_with)
                results.append({
                    "type": "paragraph",
                    "index": i,
                    "original": original_text,
                    "replaced": replaced_text,
                    "count": original_text.count(keyword)
                })
                
                # If not in preview mode, perform replacement
                if not preview_only:
                    paragraph.text = replaced_text
        
        # Search in tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if keyword in cell.text:
                        # Save original text and replaced text
                        original_text = cell.text
                        replaced_text = original_text.replace(keyword, replace_with)
                        results.append({
                            "type": "table cell",
                            "table_index": t_idx,
                            "row": r_idx,
                            "column": c_idx,
                            "original": original_text,
                            "replaced": replaced_text,
                            "count": original_text.count(keyword)
                        })
                        
                        # If not in preview mode, perform replacement
                        if not preview_only:
                            # Replace all paragraphs in the cell with the replaced text
                            for para in cell.paragraphs:
                                if keyword in para.text:
                                    para.text = para.text.replace(keyword, replace_with)
        
        if not results:
            return f"Keyword '{keyword}' not found"
        
        # Calculate total replacements
        total_replacements = sum(item["count"] for item in results)
        
        # Build response
        action_word = "Preview" if preview_only else "Replace"
        response = f"{action_word} '{keyword}' with '{replace_with}', found {len(results)} locations, {total_replacements} occurrences:\n\n"
        
        for idx, result in enumerate(results):
            response += f"{idx+1}. In {result['type']} "
            if result['type'] == "paragraph":
                response += f"index {result['index']} {action_word.lower()}ing {result['count']} times:\n"
            else:
                response += f"table {result['table_index']} at cell ({result['row']},{result['column']}) {action_word.lower()}ing {result['count']} times:\n"
            
            # Display original and replaced text snippets (context)
            max_display = 50
            if len(result['original']) > max_display * 2:
                # Find keyword position and display surrounding text
                start_pos = result['original'].find(keyword)
                start_pos = max(0, start_pos - max_display)
                excerpt_original = "..." + result['original'][start_pos:start_pos + max_display * 2] + "..."
                excerpt_replaced = "..." + result['replaced'][start_pos:start_pos + max_display * 2] + "..."
            else:
                excerpt_original = result['original']
                excerpt_replaced = result['replaced']
            
            response += f"  Original: {excerpt_original}\n"
            response += f"  Replaced: {excerpt_replaced}\n\n"
        
        if preview_only:
            response += "This is a preview of replacements. No actual changes were made. To execute replacements, set preview_only to False."
        else:
            response += "Replacements completed successfully."
        
        return response
    except Exception as e:
        error_msg = f"Search and replace failed: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def find_and_replace(ctx: Context, find_text: str, replace_text: str) -> str:
    """
    Find and replace text in the document

    Parameters:
    - find_text: Text to find
    - replace_text: Text to replace with
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        replace_count = 0
        
        # Find and replace in paragraphs
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                paragraph.text = paragraph.text.replace(find_text, replace_text)
                replace_count += paragraph.text.count(replace_text)
        
        # Find and replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if find_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(find_text, replace_text)
                            replace_count += paragraph.text.count(replace_text)
        
        return f"Replaced '{find_text}' with '{replace_text}', {replace_count} occurrences"
    except Exception as e:
        error_msg = f"Find and replace failed: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def merge_table_cells(
    ctx: Context,
    table_index: int,
    start_row: int,
    start_col: int,
    end_row: int,
    end_col: int
) -> str:
    """
    Merge table cells

    Parameters:
    - table_index: Table index
    - start_row: Start row index
    - start_col: Start column index
    - end_row: End row index
    - end_col: End column index
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        if not doc.tables:
            return "No tables in document"
        
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        
        table = doc.tables[table_index]
        
        # Check if row and column indices are valid
        if start_row < 0 or start_row >= len(table.rows):
            return f"Start row index out of range: {start_row}, table has {len(table.rows)} rows"
        
        if start_col < 0 or start_col >= len(table.columns):
            return f"Start column index out of range: {start_col}, table has {len(table.columns)} columns"
        
        if end_row < start_row or end_row >= len(table.rows):
            return f"End row index invalid: {end_row}, should be between {start_row} and {len(table.rows)-1}"
        
        if end_col < start_col or end_col >= len(table.columns):
            return f"End column index invalid: {end_col}, should be between {start_col} and {len(table.columns)-1}"
        
        # Get start and end cells
        start_cell = table.cell(start_row, start_col)
        end_cell = table.cell(end_row, end_col)
        
        # Merge cells
        start_cell.merge(end_cell)
        
        return f"Merged cells in table {table_index} from ({start_row},{start_col}) to ({end_row},{end_col})"
    except Exception as e:
        error_msg = f"Failed to merge table cells: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def split_table(ctx: Context, table_index: int, row_index: int) -> str:
    """
    Split table into two tables at specified row

    Parameters:
    - table_index: Table index
    - row_index: Split table after this row
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        if not doc.tables:
            return "No tables in document"
        
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        
        table = doc.tables[table_index]
        
        if row_index < 0 or row_index >= len(table.rows) - 1:
            return f"Row index invalid: {row_index}, should be between 0 and {len(table.rows)-2}"
        
        # Use XML operations to split table
        # Get table element
        tbl = table._tbl
        
        # Calculate split position
        split_position = row_index + 1
        
        # Create new table element
        new_tbl = OxmlElement('w:tbl')
        
        # Copy table properties
        for child in tbl.xpath('./w:tblPr')[0].getchildren():
            new_tbl.append(child.copy())
        
        # Copy table grid settings
        for child in tbl.xpath('./w:tblGrid')[0].getchildren():
            new_tbl.append(child.copy())
        
        # Move rows to new table
        rows = tbl.xpath('./w:tr')
        for i in range(split_position, len(rows)):
            new_tbl.append(rows[i])
        
        # Insert new table after original table
        tbl.addnext(new_tbl)
        
        return f"Split table {table_index} after row {row_index}"
    except Exception as e:
        error_msg = f"Failed to split table: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def add_table_row(ctx: Context, table_index: int, data: Optional[list] = None) -> str:
    """
    Add a row to table

    Parameters:
    - table_index: Table index
    - data: Row data in list format
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        if not doc.tables:
            return "No tables in document"
        
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        
        table = doc.tables[table_index]
        
        # Add new row
        new_row = table.add_row()
        
        # Fill row data
        if data:
            for i, cell_text in enumerate(data):
                if i < len(new_row.cells):
                    new_row.cells[i].text = str(cell_text)
        
        return f"Added new row to table {table_index}"
    except Exception as e:
        error_msg = f"Failed to add table row: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def delete_table_row(ctx: Context, table_index: int, row_index: int) -> str:
    """
    Delete a row from table

    Parameters:
    - table_index: Table index
    - row_index: Row index to delete
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        if not doc.tables:
            return "No tables in document"
        
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        
        table = doc.tables[table_index]
        
        if row_index < 0 or row_index >= len(table.rows):
            return f"Row index out of range: {row_index}, table has {len(table.rows)} rows"
        
        # Use XML operations to delete row
        row = table.rows[row_index]._tr
        row.getparent().remove(row)
        
        return f"Deleted row {row_index} from table {table_index}"
    except Exception as e:
        error_msg = f"Failed to delete table row: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def edit_table_cell(ctx: Context, table_index: int, row_index: int, col_index: int, text: str) -> str:
    """
    Edit table cell content

    Parameters:
    - table_index: Table index
    - row_index: Row index
    - col_index: Column index
    - text: Cell text
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        if not doc.tables:
            return "No tables in document"
        
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"
        
        table = doc.tables[table_index]
        
        if row_index < 0 or row_index >= len(table.rows):
            return f"Row index out of range: {row_index}, table has {len(table.rows)} rows"
        
        if col_index < 0 or col_index >= len(table.columns):
            return f"Column index out of range: {col_index}, table has {len(table.columns)} columns"
        
        # Modify cell content
        table.cell(row_index, col_index).text = text
        
        return f"Cell ({row_index}, {col_index}) in table {table_index} has been modified"
    except Exception as e:
        error_msg = f"Failed to edit table cell: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def add_page_break(ctx: Context) -> str:
    """
    Add page break
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        processor.current_document.add_page_break()
        
        return "Page break added"
    except Exception as e:
        error_msg = f"Failed to add page break: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def set_page_margins(
    ctx: Context,
    top: Optional[float] = None,
    bottom: Optional[float] = None,
    left: Optional[float] = None,
    right: Optional[float] = None
) -> str:
    """
    Set page margins

    Parameters:
    - top: Top margin (cm)
    - bottom: Bottom margin (cm)
    - left: Left margin (cm)
    - right: Right margin (cm)
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        # Get current section (default to use first section)
        section = doc.sections[0]
        
        # Set page margins
        if top is not None:
            section.top_margin = Cm(top)
        if bottom is not None:
            section.bottom_margin = Cm(bottom)
        if left is not None:
            section.left_margin = Cm(left)
        if right is not None:
            section.right_margin = Cm(right)
        
        return "Page margins set"
    except Exception as e:
        error_msg = f"Failed to set page margins: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def delete_paragraph(ctx: Context, paragraph_index: int) -> str:
    """
    Delete specified paragraph from document

    Parameters:
    - paragraph_index: Paragraph index to delete
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}, document has {len(doc.paragraphs)} paragraphs"
        
        # python-docx does not provide a direct method to delete a paragraph, use XML operations
        paragraph = doc.paragraphs[paragraph_index]
        p = paragraph._element
        p.getparent().remove(p)
        # Delete paragraph object reference for garbage collection
        paragraph._p = None
        paragraph._element = None
        
        return f"Paragraph {paragraph_index} deleted"
    except Exception as e:
        error_msg = f"Failed to delete paragraph: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def delete_text(ctx: Context, paragraph_index: int, start_pos: int, end_pos: int) -> str:
    """
    Delete specified text from paragraph

    Parameters:
    - paragraph_index: Paragraph index
    - start_pos: Start position (0-based index)
    - end_pos: End position (not included in the text)
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}, document has {len(doc.paragraphs)} paragraphs"
        
        paragraph = doc.paragraphs[paragraph_index]
        text = paragraph.text
        
        if start_pos < 0 or start_pos >= len(text):
            return f"Start position out of range: {start_pos}, paragraph length is {len(text)}"
        
        if end_pos <= start_pos or end_pos > len(text):
            return f"End position invalid: {end_pos}, should be between {start_pos+1} and {len(text)}"
        
        # Build new text (delete specified text)
        new_text = text[:start_pos] + text[end_pos:]
        paragraph.text = new_text
        
        return f"Deleted text from position {start_pos} to {end_pos} in paragraph {paragraph_index}"
    except Exception as e:
        error_msg = f"Failed to delete text: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def save_as_document(ctx: Context, new_file_path: str) -> str:
    """
    Save current document as a new file

    Parameters:
    - new_file_path: Path to save the new file
    """
    try:
        if err := _check_readonly():
            return err
        if err := _check_path_allowed(new_file_path):
            return err
        if not processor.current_document:
            return "No document is open"
        
        # Save as new file
        processor.current_document.save(new_file_path)
        
        # Update current file path
        processor.current_file_path = new_file_path
        processor.documents[new_file_path] = processor.current_document
        
        return f"Document saved as: {new_file_path}"
    except Exception as e:
        error_msg = f"Failed to save document: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def create_document_copy(ctx: Context, suffix: str = "-副本") -> str:
    """
    Create a copy of the current document in the directory of the original file

    Parameters:
    - suffix: Suffix to add to the original file name, default is "-副本"
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        if not processor.current_file_path:
            return "Current document has not been saved, cannot create a copy"
        
        # Parse original file path
        file_dir = os.path.dirname(processor.current_file_path)
        file_name = os.path.basename(processor.current_file_path)
        file_name_without_ext, file_ext = os.path.splitext(file_name)
        
        # Create new file name
        new_file_name = f"{file_name_without_ext}{suffix}{file_ext}"
        new_file_path = os.path.join(file_dir, new_file_name)
        
        # Save as new file
        processor.current_document.save(new_file_path)
        
        return f"Document copy created: {new_file_path}"
    except Exception as e:
        error_msg = f"Failed to create document copy: {str(e)}"
        logger.error(error_msg)
        return error_msg

@mcp.tool()
def replace_section(ctx: Context, section_title: str, new_content: list, preserve_title: bool = True) -> str:
    """
    Find specified title in document and replace content under that title, keeping original position, format, and style

    Parameters:
    - section_title: Title text to find
    - new_content: New content list, each element is a paragraph
    - preserve_title: Whether to keep original title, default is True
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        # Find title position
        title_index = -1
        for i, paragraph in enumerate(doc.paragraphs):
            if section_title in paragraph.text:
                title_index = i
                break
        
        if title_index == -1:
            return f"Title not found: '{section_title}'"
        
        # Determine end position of that section (next same or higher level title)
        end_index = len(doc.paragraphs)
        title_style = doc.paragraphs[title_index].style
        
        for i in range(title_index + 1, len(doc.paragraphs)):
            # If next same level or higher level title found, set as end position
            if doc.paragraphs[i].style.name.startswith('Heading') and \
               (doc.paragraphs[i].style.name <= title_style.name or doc.paragraphs[i].style == title_style):
                end_index = i
                break
        
        # Save original paragraph style and format information
        original_styles = []
        for i in range(start_delete := (title_index + (1 if preserve_title else 0)), min(end_index, start_delete + len(new_content))):
            if i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                style_info = {
                    'style': para.style,
                    'alignment': para.alignment,
                    'runs': []
                }
                
                # Save each run format
                for run in para.runs:
                    run_info = {
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_size': run.font.size,
                        'font_name': run.font.name,
                        'color': run.font.color.rgb if run.font.color.rgb else None
                    }
                    style_info['runs'].append(run_info)
                
                original_styles.append(style_info)
            else:
                # If original paragraph count is insufficient, use last paragraph style
                if original_styles:
                    original_styles.append(original_styles[-1])
                else:
                    # If no original style, use default style
                    original_styles.append({
                        'style': None,
                        'alignment': None,
                        'runs': []
                    })
        
        # If original style count is insufficient, use last style to fill
        while len(original_styles) < len(new_content):
            if original_styles:
                original_styles.append(original_styles[-1])
            else:
                original_styles.append({
                    'style': None,
                    'alignment': None,
                    'runs': []
                })
        
        # Record insert position
        insert_position = start_delete
        
        # Delete from end to avoid index change
        for i in range(end_index - 1, start_delete - 1, -1):
            p = doc.paragraphs[i]._element
            p.getparent().remove(p)
        
        # Add new content, apply original format
        for i, content in enumerate(reversed(new_content)):
            # Create new paragraph
            p = doc.add_paragraph()
            
            # Apply original paragraph style
            style_info = original_styles[len(new_content) - i - 1]
            if style_info['style']:
                p.style = style_info['style']
            if style_info['alignment'] is not None:
                p.alignment = style_info['alignment']
            
            # Add text and apply format
            if style_info['runs'] and len(style_info['runs']) > 0:
                # If multiple runs, try to keep format
                # Simplified processing: Add entire content to a run, apply format from first run
                run = p.add_run(content)
                run_info = style_info['runs'][0]
                
                run.bold = run_info['bold']
                run.italic = run_info['italic']
                run.underline = run_info['underline']
                
                if run_info['font_size']:
                    run.font.size = run_info['font_size']
                
                if run_info['font_name']:
                    run.font.name = run_info['font_name']
                    # Set Chinese font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), run_info['font_name'])
                
                if run_info['color']:
                    run.font.color.rgb = run_info['color']
            else:
                # If no run information, add text directly
                p.text = content
            
            # Move new paragraph to correct position
            doc._body._body.insert(insert_position, p._p)
            
            # Delete original added paragraph (at end of document)
            doc._body._body.remove(doc.paragraphs[-1]._p)
        
        return f"Replaced content under title '{section_title}', keeping original format and style"
    except Exception as e:
        error_msg = f"Failed to replace content: {str(e)}"
        logger.error(error_msg)
        traceback.print_exc()  # Print detailed error information
        return error_msg

@mcp.tool()
def edit_section_by_keyword(ctx: Context, keyword: str, new_content: list, section_range: int = 3) -> str:
    """
    Find paragraphs containing specified keyword and replace them and their surrounding content, keeping original position, format, and style

    Parameters:
    - keyword: Keyword to find
    - new_content: New content list, each element is a paragraph
    - section_range: Surrounding paragraph range to replace, default is 3
    """
    try:
        if err := _check_readonly():
            return err
        if not processor.current_document:
            return "No document is open"
        
        doc = processor.current_document
        
        # Find keyword position
        keyword_indices = []
        for i, paragraph in enumerate(doc.paragraphs):
            if keyword in paragraph.text:
                keyword_indices.append(i)
        
        if not keyword_indices:
            return f"Keyword not found: '{keyword}'"
        
        # Use first match
        keyword_index = keyword_indices[0]
        
        # Determine paragraph range to replace
        start_index = max(0, keyword_index - section_range)
        end_index = min(len(doc.paragraphs), keyword_index + section_range + 1)
        
        # Save original paragraph style and format information
        original_styles = []
        for i in range(start_index, min(end_index, start_index + len(new_content))):
            if i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                style_info = {
                    'style': para.style,
                    'alignment': para.alignment,
                    'runs': []
                }
                
                # Save each run format
                for run in para.runs:
                    run_info = {
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_size': run.font.size,
                        'font_name': run.font.name,
                        'color': run.font.color.rgb if run.font.color.rgb else None
                    }
                    style_info['runs'].append(run_info)
                
                original_styles.append(style_info)
            else:
                # If original paragraph count is insufficient, use last paragraph style
                if original_styles:
                    original_styles.append(original_styles[-1])
                else:
                    # If no original style, use default style
                    original_styles.append({
                        'style': None,
                        'alignment': None,
                        'runs': []
                    })
        
        # If original style count is insufficient, use last style to fill
        while len(original_styles) < len(new_content):
            if original_styles:
                original_styles.append(original_styles[-1])
            else:
                original_styles.append({
                    'style': None,
                    'alignment': None,
                    'runs': []
                })
        
        # Record insert position
        insert_position = start_index
        
        # Delete from end to avoid index change
        for i in range(end_index - 1, start_index - 1, -1):
            p = doc.paragraphs[i]._element
            p.getparent().remove(p)
        
        # Add new content, apply original format
        for i, content in enumerate(reversed(new_content)):
            # Create new paragraph
            p = doc.add_paragraph()
            
            # Apply original paragraph style
            style_info = original_styles[len(new_content) - i - 1]
            if style_info['style']:
                p.style = style_info['style']
            if style_info['alignment'] is not None:
                p.alignment = style_info['alignment']
            
            # Add text and apply format
            if style_info['runs'] and len(style_info['runs']) > 0:
                # If multiple runs, try to keep format
                # Simplified processing: Add entire content to a run, apply format from first run
                run = p.add_run(content)
                run_info = style_info['runs'][0]
                
                run.bold = run_info['bold']
                run.italic = run_info['italic']
                run.underline = run_info['underline']
                
                if run_info['font_size']:
                    run.font.size = run_info['font_size']
                
                if run_info['font_name']:
                    run.font.name = run_info['font_name']
                    # Set Chinese font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), run_info['font_name'])
                
                if run_info['color']:
                    run.font.color.rgb = run_info['color']
            else:
                # If no run information, add text directly
                p.text = content
            
            # Move new paragraph to correct position
            doc._body._body.insert(insert_position, p._p)
            
            # Delete original added paragraph (at end of document)
            doc._body._body.remove(doc.paragraphs[-1]._p)
        
        return f"Replaced paragraphs containing keyword '{keyword}' and their surrounding content, keeping original format and style"
    except Exception as e:
        error_msg = f"Failed to replace content: {str(e)}"
        logger.error(error_msg)
        traceback.print_exc()  # Print detailed error information
        return error_msg

@mcp.tool()
def get_headings(ctx: Context) -> str:
    """
    Get all headings in the document with their level and position index.
    Useful for understanding document structure before reading specific sections.
    """
    try:
        if not processor.current_document:
            return "No document is open"

        doc = processor.current_document
        headings = []

        for i, paragraph in enumerate(doc.paragraphs):
            style_name = paragraph.style.name
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split(" ")[-1])
                except ValueError:
                    level = 0
                headings.append({
                    "index": i,
                    "level": level,
                    "text": paragraph.text.strip()
                })

        if not headings:
            return "No headings found in document"

        lines = [f"Found {len(headings)} headings:\n"]
        for h in headings:
            indent = "  " * (h["level"] - 1)
            lines.append(f"{indent}[H{h['level']}] (paragraph {h['index']}) {h['text']}")

        return "\n".join(lines)
    except Exception as e:
        error_msg = f"Failed to get headings: {str(e)}"
        logger.error(error_msg)
        return error_msg


@mcp.tool()
def get_full_text(ctx: Context, include_tables: bool = True) -> str:
    """
    Get the full text content of the document.

    Parameters:
    - include_tables: Whether to include table cell content (default True)
    """
    try:
        if not processor.current_document:
            return "No document is open"

        doc = processor.current_document
        parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        if include_tables:
            for t_idx, table in enumerate(doc.tables):
                parts.append(f"\n[Table {t_idx}]")
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    parts.append(" | ".join(row_cells))

        if not parts:
            return "Document is empty"

        return "\n".join(parts)
    except Exception as e:
        error_msg = f"Failed to get full text: {str(e)}"
        logger.error(error_msg)
        return error_msg


@mcp.tool()
def get_paragraph(ctx: Context, paragraph_index: int) -> str:
    """
    Get content and formatting details of a specific paragraph.

    Parameters:
    - paragraph_index: Index of the paragraph to retrieve
    """
    try:
        if not processor.current_document:
            return "No document is open"

        doc = processor.current_document

        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index out of range: {paragraph_index}, document has {len(doc.paragraphs)} paragraphs"

        para = doc.paragraphs[paragraph_index]
        info = [f"Paragraph {paragraph_index}:"]
        info.append(f"  Style: {para.style.name}")
        info.append(f"  Alignment: {para.alignment}")
        info.append(f"  Text: {para.text}")
        info.append(f"  Runs ({len(para.runs)}):")

        for r_idx, run in enumerate(para.runs):
            run_details = []
            if run.bold:
                run_details.append("bold")
            if run.italic:
                run_details.append("italic")
            if run.underline:
                run_details.append("underline")
            if run.font.size:
                run_details.append(f"{run.font.size.pt:.0f}pt")
            if run.font.name:
                run_details.append(f"font={run.font.name}")
            try:
                if run.font.color.rgb:
                    run_details.append(f"color=#{run.font.color.rgb}")
            except Exception:
                pass
            fmt = ", ".join(run_details) if run_details else "default"
            info.append(f"    Run {r_idx}: [{fmt}] '{run.text}'")

        return "\n".join(info)
    except Exception as e:
        error_msg = f"Failed to get paragraph: {str(e)}"
        logger.error(error_msg)
        return error_msg


@mcp.tool()
def get_paragraphs(ctx: Context, start_index: int = 0, end_index: Optional[int] = None) -> str:
    """
    Get a range of paragraphs from the document.

    Parameters:
    - start_index: First paragraph index (inclusive, default 0)
    - end_index: Last paragraph index (exclusive, default is end of document)
    """
    try:
        if not processor.current_document:
            return "No document is open"

        doc = processor.current_document
        total = len(doc.paragraphs)

        if end_index is None:
            end_index = total

        start_index = max(0, start_index)
        end_index = min(total, end_index)

        if start_index >= end_index:
            return f"Invalid range: [{start_index}, {end_index}), document has {total} paragraphs"

        lines = [f"Paragraphs {start_index} to {end_index - 1} (of {total} total):\n"]
        for i in range(start_index, end_index):
            para = doc.paragraphs[i]
            style = para.style.name
            text = para.text if para.text.strip() else "(empty)"
            lines.append(f"[{i}] ({style}) {text}")

        return "\n".join(lines)
    except Exception as e:
        error_msg = f"Failed to get paragraphs: {str(e)}"
        logger.error(error_msg)
        return error_msg


@mcp.tool()
def get_table_content(ctx: Context, table_index: int) -> str:
    """
    Get the full content of a specific table, including all rows and cells.

    Parameters:
    - table_index: Index of the table to retrieve
    """
    try:
        if not processor.current_document:
            return "No document is open"

        doc = processor.current_document

        if not doc.tables:
            return "No tables in document"

        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index out of range: {table_index}, document has {len(doc.tables)} tables"

        table = doc.tables[table_index]
        rows_count = len(table.rows)
        cols_count = len(table.columns)

        lines = [f"Table {table_index} ({rows_count} rows x {cols_count} columns):\n"]

        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(f"  Row {r_idx}: {' | '.join(cells)}")

        return "\n".join(lines)
    except Exception as e:
        error_msg = f"Failed to get table content: {str(e)}"
        logger.error(error_msg)
        return error_msg


@mcp.tool()
def get_section_content(ctx: Context, section_title: str) -> str:
    """
    Get all content under a specific heading/section title.
    Returns all paragraphs from the heading until the next heading of same or higher level.

    Parameters:
    - section_title: The heading text to find (partial match is supported)
    """
    try:
        if not processor.current_document:
            return "No document is open"

        doc = processor.current_document
        paragraphs = doc.paragraphs

        # Find the heading
        title_index = -1
        for i, para in enumerate(paragraphs):
            if section_title.lower() in para.text.lower():
                title_index = i
                break

        if title_index == -1:
            return f"Section title not found: '{section_title}'"

        title_para = paragraphs[title_index]
        title_style = title_para.style.name

        # Determine title level
        try:
            title_level = int(title_style.split(" ")[-1]) if title_style.startswith("Heading") else 0
        except ValueError:
            title_level = 0

        # Find end of section: next heading of same or higher level
        end_index = len(paragraphs)
        for i in range(title_index + 1, len(paragraphs)):
            style = paragraphs[i].style.name
            if style.startswith("Heading"):
                try:
                    level = int(style.split(" ")[-1])
                except ValueError:
                    level = 0
                if title_level == 0 or level <= title_level:
                    end_index = i
                    break

        lines = [f"Section: '{title_para.text.strip()}' (paragraphs {title_index}-{end_index - 1})\n"]
        for i in range(title_index + 1, end_index):
            text = paragraphs[i].text
            if text.strip():
                lines.append(f"[{i}] {text}")

        if len(lines) == 1:
            lines.append("(section has no body content)")

        return "\n".join(lines)
    except Exception as e:
        error_msg = f"Failed to get section content: {str(e)}"
        logger.error(error_msg)
        return error_msg


@mcp.tool()
def get_document_outline(ctx: Context) -> str:
    """
    Get a hierarchical outline of the document showing all headings and paragraph counts per section.
    """
    try:
        if not processor.current_document:
            return "No document is open"

        doc = processor.current_document
        paragraphs = doc.paragraphs

        outline = []
        current_section_para_count = 0
        current_heading = None

        for i, para in enumerate(paragraphs):
            style = para.style.name
            if style.startswith("Heading"):
                # Record previous section paragraph count
                if current_heading is not None:
                    current_heading["paragraphs"] = current_section_para_count
                    outline.append(current_heading)

                try:
                    level = int(style.split(" ")[-1])
                except ValueError:
                    level = 0

                current_heading = {
                    "index": i,
                    "level": level,
                    "text": para.text.strip(),
                    "paragraphs": 0
                }
                current_section_para_count = 0
            elif para.text.strip():
                current_section_para_count += 1

        # Add the last section
        if current_heading is not None:
            current_heading["paragraphs"] = current_section_para_count
            outline.append(current_heading)

        if not outline:
            non_empty = sum(1 for p in paragraphs if p.text.strip())
            return f"Document has no headings. Total non-empty paragraphs: {non_empty}"

        lines = [f"Document outline ({len(outline)} sections):\n"]
        for section in outline:
            indent = "  " * (section["level"] - 1)
            para_info = f"({section['paragraphs']} paragraphs)" if section["paragraphs"] > 0 else "(no body)"
            lines.append(f"{indent}[H{section['level']}] {section['text']} {para_info}")

        total_tables = len(doc.tables)
        if total_tables:
            lines.append(f"\nDocument also contains {total_tables} table(s).")

        return "\n".join(lines)
    except Exception as e:
        error_msg = f"Failed to get document outline: {str(e)}"
        logger.error(error_msg)
        return error_msg


# Add more tools...

if __name__ == "__main__":
    if READONLY_MODE:
        logger.info("Server starting in READ-ONLY mode: write operations are disabled")
    if CWD_ONLY_MODE:
        logger.info(f"Server starting in CWD-ONLY mode: file access restricted to {os.getcwd()}")

    # Always start with a clean state, don't try to load any previous document
    if os.path.exists(CURRENT_DOC_FILE):
        try:
            os.remove(CURRENT_DOC_FILE)
            logger.info("Removed existing state file for clean startup")
        except Exception as e:
            logger.error(f"Failed to remove existing state file: {e}")
    
    # Run MCP server
    mcp.run() 